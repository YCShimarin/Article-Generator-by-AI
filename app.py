from fileinput import filename
import os, json, requests, time, re, random
from datetime import datetime
from pathlib import Path
import markdown
from xhtml2pdf import pisa
from docx import Document
from io import BytesIO
from flask import Flask, render_template, request, jsonify, send_from_directory, Response, stream_with_context, send_file

app = Flask(__name__)

# --- KONFIGURASI ---
OUTPUT_DIR = "output"
LOG_DIR = "logs"
PROMPT_DIR = "prompts"
FALLBACK_MODEL = "stepfun/step-3.5-flash:free"
URL = "https://openrouter.ai/api/v1/chat/completions"

# Pastikan folder ada
for d in [OUTPUT_DIR, LOG_DIR, PROMPT_DIR]:
    os.makedirs(d, exist_ok=True)

# --- UTILS (DIADAPTASI DARI KODEMU) ---
AVAILABLE_MODES = {
    "science":      "🔬 Science / Academic (with LaTeX)",
    "sci_fi_story": "🚀 Science Fiction Story",
    "fiction_story": "📖 Fiction Story",
    "blog_tech":    "💻 Tech Blog Article",
    "history":      "🏛️ History Article",
}

def load_prompt_file(name, default="", mode=None):
    # Try mode-specific subfolder first
    if mode and mode in AVAILABLE_MODES:
        mode_path = os.path.join(PROMPT_DIR, mode, name)
        if os.path.exists(mode_path):
            with open(mode_path, "r", encoding="utf-8") as f: return f.read()
    # Fallback to root prompts folder
    path = os.path.join(PROMPT_DIR, name)
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f: return f.read()
    return default

def sanitize_filename(text):
    text = text.lower().strip()
    text = re.sub(r"[\\/*?:\"<>|]", "", text)
    text = re.sub(r"\s+", "-", text)
    return text

def count_words(text):
    return len(re.findall(r"\b\w+\b", text, flags=re.UNICODE))

# --- CORE API CALL (ROBUST VERSION) ---
def call_openrouter(token, model, prompt, retries=3):
    current_model = model
    for attempt in range(retries):
        try:
            payload = {
                "model": current_model,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.7,
                "max_tokens": 1200
            }
            res = requests.post(URL, headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"}, json=payload, timeout=40)
            data = res.json()
            
            if "choices" in data:
                return data["choices"][0]["message"]["content"].strip()
            
            if "error" in data and data["error"].get("code") == 429:
                time.sleep(10) # Rate limit hit
            
            if attempt == retries - 2: # Last attempt, switch to fallback
                current_model = FALLBACK_MODEL
        except:
            time.sleep(2)
    return None

def generate_seo_metadata(token, model, topic, content, lang="English"):
    """Generates SEO focus keywords, meta description, and slug."""
    prompt = f"""
    Based on this article topic: '{topic}' 
    And some content: {content[:500]}...
    
    Generate SEO metadata in {lang}:
    - Focus Keywords (3-5)
    - Meta Description (max 160 chars)
    - Suggested URL Slug
    
    Format the output EXACTLY like this (Markdown):
    <details>
    <summary>🔍 SEO_OPTIMIZATION_METADATA (Click to Expand)</summary>
    
    **Focus Keywords:** [keyword1, keyword2...]
    **Meta Description:** [description...]
    **Suggested Slug:** [slug...]
    
    </details>
    """
    return call_openrouter(token, model, prompt) or ""

# --- FLASK ROUTES ---

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/prompt-modes')
def prompt_modes():
    """Returns the list of available prompt modes for the frontend dropdown."""
    modes = [{"id": k, "label": v} for k, v in AVAILABLE_MODES.items()]
    return jsonify({"modes": modes})

@app.route('/suggest-topics', methods=['POST'])
def suggest_topics():
    data = request.json
    token = data.get('token')
    model = data.get('model')
    keyword = data.get('keyword')
    lang = data.get('lang', 'English')
    
    prompt = f"Suggest 3 catchy and professional article titles about '{keyword}' in {lang}. Output only the titles, one per line, no numbering, no explanations."
    
    response = call_openrouter(token, model, prompt)
    if response:
        titles = [t.strip() for t in response.split('\n') if t.strip()][:3]
        return jsonify({"status": "success", "titles": titles})
    return jsonify({"status": "error", "message": "Failed to suggest topics"}), 500

@app.route('/export/<format>/<path:filename>')
def export_article(format, filename):
    filepath = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(filepath):
        return "File not found", 404
        
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    
    # Remove metadata if present (the part starting with >)
    clean_content = re.sub(r'^>.*?\n\n', '', content, flags=re.DOTALL)
    
    if format == 'md':
        return send_file(filepath, as_attachment=True)
        
    elif format == 'pdf':
        # Convert Markdown to HTML
        html_content = markdown.markdown(clean_content)
        # Add basic styling for PDF
        styled_html = f"<html><head><style>body {{ font-family: Helvetica, Arial, sans-serif; line-height: 1.6; }} h1 {{ color: #333; }}</style></head><body>{html_content}</body></html>"
        
        pdf_buffer = BytesIO()
        pisa_status = pisa.CreatePDF(styled_html, dest=pdf_buffer)
        
        if pisa_status.err:
            return "Error generating PDF", 500
            
        pdf_buffer.seek(0)
        return send_file(pdf_buffer, mimetype='application/pdf', as_attachment=True, download_name=filename.replace('.md', '.pdf'))
        
    elif format == 'docx':
        doc = Document()
        # Very basic conversion: title and then paragraphs
        lines = clean_content.split('\n')
        for line in lines:
            if line.startswith('# '):
                doc.add_heading(line[2:], level=0)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=1)
            elif line.strip():
                doc.add_paragraph(line)
        
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return send_file(docx_buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name=filename.replace('.md', '.docx'))

    return "Invalid format", 400

@app.route('/generate-articles', methods=['POST'])
def generate_articles():
    data = request.json
    token = data.get('token')
    model = data.get('model')
    titles = data.get('titles')
    mode = data.get('mode', 'science')  # Default to science
    lang = data.get('lang', 'English') # Default to English

    # Load Templates (mode-aware)
    outline_temp = load_prompt_file("outline_prompt.txt", "Create an outline for {TOPIC}", mode=mode)
    section_temp = load_prompt_file("section_prompt.txt", "Write section {SECTION} for {TOPIC}", mode=mode)
    style = load_prompt_file("style.txt", "Technical, clear, professional.", mode=mode)
    
    # Inject language
    outline_temp = outline_temp.replace("{LANGUAGE}", lang)
    section_temp = section_temp.replace("{LANGUAGE}", lang)

    generated_count = 0

    for topic in titles:
        # STEP 1: GENERATE OUTLINE
        prompt_outline = outline_temp.replace("{TOPIC}", topic)
        outline_raw = call_openrouter(token, model, prompt_outline)
        
        if not outline_raw: continue

        # FORCE CONTROL: Parse & Limit Sections
        # Kita ambil baris yang terlihat seperti list (1. Judul atau - Judul)
        lines = outline_raw.split('\n')
        sections = []
        for line in lines:
            clean = re.sub(r'^\d+\.\s*|^\-\s*', '', line).strip()
            if clean and len(clean) > 5: # Hindari baris sampah
                sections.append(clean)
        
        # --- THE GUARDRAIL ---
        if len(sections) > 8:
            sections = sections[:8] # Potong paksa jika AI labil
        
        # STEP 2: GENERATE SECTIONS
        article_body = f"# {topic}\n\n"
        start_time = time.time()
        
        for sec_title in sections:
            prompt_sec = section_temp.replace("{TOPIC}", topic).replace("{SECTION}", sec_title).replace("{STYLE}", style)
            content = call_openrouter(token, model, prompt_sec)
            if content:
                article_body += f"## {sec_title}\n\n{content}\n\n"
                time.sleep(random.uniform(2, 4)) # Delay antar section agar aman

        # STEP 2.5: GENERATE SEO METADATA
        seo_meta = generate_seo_metadata(token, model, topic, article_body, lang=lang)
        final_md_with_seo = article_body + "\n\n---\n" + seo_meta

        # STEP 3: METADATA & SAVE (PERSIS GAYA KODEMU)
        total_time = time.time() - start_time
        words = count_words(article_body)
        metadata = f"> **AI Metadata**\n> - Model: {model}\n> - Time: {int(total_time)}s\n> - Words: {words}\n> - Sections: {len(sections)}\n\n"
        
        final_md = metadata + final_md_with_seo
        filename = sanitize_filename(topic) + ".md"
        with open(os.path.join(OUTPUT_DIR, filename), "w", encoding="utf-8") as f:
            f.write(final_md)
        
        generated_count += 1

    # Update JSON Index Otomatis
    sync_files = [f for f in os.listdir(OUTPUT_DIR) if f.endswith(".md")]
    with open(os.path.join(OUTPUT_DIR, "articles.json"), "w") as f:
        json.dump({"articles": sync_files}, f, indent=2)

    return jsonify({"status": "success", "count": generated_count})

@app.route('/sync-archive', methods=['POST'])
def sync_archive():
    files = [f for f in os.listdir(OUTPUT_DIR) if f.endswith(".md")]
    with open(os.path.join(OUTPUT_DIR, "articles.json"), "w") as f:
        json.dump({"articles": files}, f, indent=2)
    return jsonify({"status": "success", "count": len(files)})

@app.route('/output/<path:filename>')
def serve_output(filename):
    return send_from_directory(OUTPUT_DIR, filename)
@app.route('/generate-stream')
def generate_stream():
    token = request.args.get('token')
    model = request.args.get('model')
    raw_titles = request.args.get('titles', '')
    mode = request.args.get('mode', 'science')  # Default to science
    lang = request.args.get('lang', 'English') # Default to English
    titles = [t.strip() for t in raw_titles.split(',') if t.strip()]

    def generate():
        start_time_all = time.time()
        current_filename = None
        
        mode_label = AVAILABLE_MODES.get(mode, mode)
        yield f"data: {json.dumps({'msg': f'🚀 ENGINE STARTED | MODE: {mode_label} | LANG: {lang}', 'type': 'success'})}\n\n"
        
        try:
            outline_temp = load_prompt_file("outline_prompt.txt", mode=mode).replace("{LANGUAGE}", lang)
            section_temp = load_prompt_file("section_prompt.txt", mode=mode).replace("{LANGUAGE}", lang)
            style_guide = load_prompt_file("style.txt", mode=mode)

            for topic in titles:
                yield f"data: {json.dumps({'msg': f'📌 WORKING ON: {topic}', 'type': 'info'})}\n\n"
                
                # STEP 1: OUTLINE
                yield f"data: {json.dumps({'msg': 'Step 1: Generating Outline...', 'type': 'warn'})}\n\n"
                outline_raw = call_openrouter(token, model, outline_temp.replace("{TOPIC}", topic))
                
                if not outline_raw:
                    yield f"data: {json.dumps({'msg': f'❌ Failed to get outline for {topic}', 'type': 'err'})}\n\n"
                    continue

                # Parse & Force Control (Max 8 Sections)
                sections = [re.sub(r'^\d+\.\s*|^\-\s*', '', line).strip() 
                            for line in outline_raw.split('\n') if len(line.strip()) > 5]
                if len(sections) > 8:
                    sections = sections[:8]
                    yield f"data: {json.dumps({'msg': '⚠️ Sections truncated to 8 (Force Control)', 'type': 'warn'})}\n\n"

                # STEP 2: SECTIONS
                full_article = f"# {topic}\n\n"
                current_filename = sanitize_filename(topic) + ".md"
                start_time_article = time.time()
                
                for i, sec in enumerate(sections, 1):
                    yield f"data: {json.dumps({'msg': f'✍️ Writing [{i}/{len(sections)}]: {sec}', 'type': 'info'})}\n\n"
                    p_sec = section_temp.replace("{TOPIC}", topic).replace("{SECTION}", sec).replace("{STYLE}", style_guide)
                    content = call_openrouter(token, model, p_sec)
                    
                    if content:
                        full_article += f"## {sec}\n\n{content}\n\n"
                    
                    # Small delay to prevent rate limit
                    time.sleep(random.uniform(1, 3))

                # STEP 2.5: SEO METADATA
                yield f"data: {json.dumps({'msg': '🔍 Finalizing: Generating SEO Meta-Tags...', 'type': 'warn'})}\n\n"
                seo_meta = generate_seo_metadata(token, model, topic, full_article, lang=lang)
                
                # STEP 3: SAVE
                duration = int(time.time() - start_time_article)
                words = count_words(full_article)
                metadata = f"> **Metadata**\n> - Model: {model}\n> - Mode: {mode_label}\n> - Time: {duration}s\n> - Words: {words}\n\n"
                
                with open(os.path.join(OUTPUT_DIR, current_filename), "w", encoding="utf-8") as f:
                    f.write(metadata + full_article + "\n\n---\n" + seo_meta)
                
                yield f"data: {json.dumps({'msg': f'💾 SAVED WITH SEO: {current_filename}', 'type': 'success'})}\n\n"
                current_filename = None # Successfully finished
                
            yield f"data: {json.dumps({'msg': '🏁 ALL TASKS FINISHED', 'type': 'success'})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
            
            # FINAL SYNC
            sync_archive()

        except GeneratorExit:
            print(f"!!! CLIENT DISCONNECTED !!! Stopping generation.")
            if current_filename:
                file_path = os.path.join(OUTPUT_DIR, current_filename)
                if os.path.exists(file_path):
                    print(f"Cleaning up partial file: {current_filename}")
                    try:
                        os.remove(file_path)
                    except:
                        pass
                sync_archive()
            return

    return Response(stream_with_context(generate()), mimetype='text/event-stream')

if __name__ == '__main__':
    # app.run(debug=True)
    app.run(host="0.0.0.0", port=7860, debug=False)